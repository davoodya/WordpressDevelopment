#!/usr/bin/env python3
"""
CSV -> Word converter (interactive input path).

این اسکریپت فایل CSV را به یک فایل Word تبدیل می‌کند و داده‌ها را در قالب یک جدول زیبا نمایش می‌دهد.

ویژگی‌ها:
- دریافت مسیر کامل (Absolute Path) فایل CSV از کاربر
- پشتیبانی از مسیرهایی که فاصله (Space) دارند
- ساخت خروجی در همان پوشه CSV
- نام خروجی: همان نام CSV با پسوند .docx
- نمایش داده‌ها در قالب جدول با قالب‌بندی حرفه‌ای
- پشتیبانی کامل از زبان فارسی
- شناسایی و حذف رکوردهای تکراری
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] Required package 'python-docx' is not installed.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)


def normalize_user_path(raw: str) -> Path:
    """Normalize input path from user (handles quotes and spaces)."""
    cleaned = raw.strip().strip('"').strip("'").strip()
    return Path(cleaned)


def setup_rtl_support(doc: Document):
    """تنظیم پشتیبانی از راست به چپ برای سند Word."""
    sections = doc.sections
    for section in sections:
        section.right_margin = Inches(1)
        section.left_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def convert_csv_to_word(input_csv: Path) -> tuple[Path, dict[str, int]]:
    """
    تبدیل فایل CSV به فایل Word با جدول قالب‌بندی شده.
    
    Args:
        input_csv: مسیر فایل CSV ورودی
    
    Returns:
        tuple: (مسیر فایل word خروجی, دیکشنری آمار)
        دیکشنری آمار شامل:
            - total_rows: تعداد کل ردیف‌های خوانده شده از CSV
            - duplicates_found: تعداد محصولات تکراری که نادیده گرفته شدند
            - unique_products: تعداد محصولات یکتا نوشته شده در Word
    """
    output_docx = input_csv.with_suffix(".docx")

    # ایجاد سند Word جدید
    doc = Document()
    
    # تنظیم پشتیبانی از راست به چپ
    setup_rtl_support(doc)
    
    # افزودن عنوان سند
    title = doc.add_heading("لیست محصولات و قیمت‌ها", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # افزودن فونت فارسی به عنوان
    for run in title.runs:
        run.font.name = 'B Nazanin'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # افزودن اطلاعات متا
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run(f"تاریخ ایجاد: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info_run.font.name = 'B Nazanin'
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(100, 100, 100)
    
    source_run = info_para.add_run(f"فایل منبع: {input_csv.name}")
    source_run.font.name = 'B Nazanin'
    source_run.font.size = Pt(10)
    source_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # فاصله خالی
    
    # خواندن داده‌ها از CSV
    rows_data = []
    headers = []
    
    # آمارگیری
    seen_products = set()
    total_rows = 0
    duplicates_found = 0
    unique_products = 0
    
    with input_csv.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        
        # خواندن هدر
        headers = next(reader, None)
        if not headers:
            headers = ["نام محصول", "قیمت"]
        
        # خواندن داده‌ها
        for row in reader:
            if not row:
                continue
            
            total_rows += 1
            
            product_name = row[0] if len(row) >= 1 else ""
            
            # پرش از ردیف‌های خالی
            if all(cell.strip() == "" for cell in row):
                continue
            
            # نرمال‌سازی نام محصول برای شناسایی تکراری
            normalized_name = product_name.strip().lower()
            
            if not normalized_name:
                continue
            
            # بررسی تکراری
            if normalized_name in seen_products:
                duplicates_found += 1
                continue
            
            seen_products.add(normalized_name)
            rows_data.append(row)
            unique_products += 1
    
    # ایجاد جدول در Word
    # تعداد ستون‌ها = تعداد هدرها
    num_cols = len(headers)
    num_rows = len(rows_data) + 1  # +1 برای ردیف هدر
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # تنظیم عرض جدول
    table.autofit = False
    table.allow_autofit = False
    
    # پر کردن هدر جدول
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header_text
        
        # قالب‌بندی سلول هدر
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.name = 'B Nazanin'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        # رنگ پس‌زمینه هدر
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "0070C0")
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # پر کردن داده‌های جدول
    for row_idx, row_data in enumerate(rows_data, start=1):
        row_cells = table.rows[row_idx].cells
        
        for col_idx, cell_value in enumerate(row_data):
            if col_idx >= num_cols:
                break
            
            cell = row_cells[col_idx]
            cell.text = str(cell_value)
            
            # قالب‌بندی سلول داده
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for run in paragraph.runs:
                run.font.name = 'B Nazanin'
                run.font.size = Pt(11)
                
                # رنگ متناوب برای ردیف‌ها
                if row_idx % 2 == 0:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), "E7E6E6")
                    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # افزودن پاورقی با آمار
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    footer_text = f"""
    ══════════════════════════════════════
    آمار تبدیل:
    • تعداد کل ردیف‌های پردازش شده: {total_rows}
    • تعداد محصولات یکتا: {unique_products}
    • تعداد محصولات تکراری حذف شده: {duplicates_found}
    ══════════════════════════════════════
    """
    
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = 'B Nazanin'
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(50, 50, 50)
    
    # ذخیره سند
    doc.save(output_docx)
    
    statistics = {
        "total_rows": total_rows,
        "duplicates_found": duplicates_found,
        "unique_products": unique_products,
    }
    
    return output_docx, statistics


def main() -> int:
    """تابع اصلی برنامه."""
    print("=" * 70)
    print("CSV به Word Converter - تبدیل‌گر CSV به Word")
    print("=" * 70)
    print()
    print("لطفاً مسیر کامل (Absolute Path) فایل CSV را وارد کنید:")
    print("مثال: H:\\Repo\\WordpressDevelopment\\Products-Price-Exporter\\test_sample.csv")
    print()
    raw_path = input("> ").strip()

    if not raw_path:
        print("[ERROR] مسیر فایل نمی‌تواند خالی باشد.", file=sys.stderr)
        return 1

    input_csv = normalize_user_path(raw_path)

    # بررسی صحت مسیر
    if not input_csv.is_absolute():
        print("[ERROR] مسیر ورودی باید یک مسیر کامل (Absolute Path) باشد.", file=sys.stderr)
        return 1

    if input_csv.suffix.lower() != ".csv":
        print("[ERROR] فایل ورودی باید پسوند .csv داشته باشد.", file=sys.stderr)
        return 1

    if not input_csv.exists() or not input_csv.is_file():
        print(f"[ERROR] فایل پیدا نشد: {input_csv}", file=sys.stderr)
        return 1

    print()
    print("[INFO] در حال پردازش...")
    
    try:
        output_docx, statistics = convert_csv_to_word(input_csv)
    except Exception as exc:
        print(f"[ERROR] تبدیل با خطا مواجه شد: {exc}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1

    print()
    print("=" * 70)
    print("[✓] تبدیل با موفقیت انجام شد!")
    print("=" * 70)
    print(f"[✓] فایل خروجی: {output_docx}")
    print()
    print("[آمار تبدیل]")
    print(f"  • تعداد کل ردیف‌های پردازش شده: {statistics['total_rows']}")
    print(f"  • تعداد محصولات یکتا نوشته شده: {statistics['unique_products']}")
    print(f"  • تعداد محصولات تکراری حذف شده: {statistics['duplicates_found']}")
    print("=" * 70)
    
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
